# backend/routes/upload_routes.py
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks
from pydantic import BaseModel
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
import os
import shutil
import uuid
from typing import List, Optional, Dict, Any
import logging
from datetime import datetime
import magic
from pathlib import Path
import json
import traceback

from models.database import get_db
from models.user import User
from models.document import Document, DocumentCreate, DocumentResponse, DocumentStatus, DocumentType
from .auth_routes import get_current_active_user
from rag_pipeline import get_rag_pipeline
from utils.chunking import TextChunker as LegalTextChunker
from utils.embeddings import EmbeddingGenerator

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

router = APIRouter()

# Configuration
UPLOAD_DIR = "uploads"
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
ALLOWED_MIME_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt"
}
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}

# Ensure upload directory exists
os.makedirs(UPLOAD_DIR, exist_ok=True)

class UploadResponse(BaseModel):
    success: bool
    message: str
    document_id: Optional[int] = None
    filename: Optional[str] = None
    file_size: Optional[int] = None
    processing_status: Optional[str] = None

class DocumentListResponse(BaseModel):
    success: bool
    documents: List[DocumentResponse]
    total_count: int
    page: int
    page_size: int

# ⚠️ Define DocumentUpdate locally
class DocumentUpdateModel(BaseModel):
    title: Optional[str]
    description: Optional[str]
    document_type: Optional[DocumentType]
    case_number: Optional[str]
    court_name: Optional[str]
    judge_name: Optional[str]
    parties_involved: Optional[List[str]]
    jurisdiction: Optional[str]
    legal_topics: Optional[List[str]]

# ------------------- FILE VALIDATION & SAVE -------------------
def validate_file_type(file: UploadFile) -> bool:
    """Validate file type using both extension and MIME type"""
    try:
        # Get file extension
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in ALLOWED_EXTENSIONS:
            logger.warning(f"Invalid file extension: {file_extension}")
            return False
        
        # Read first bytes for MIME detection
        file_content = file.file.read(2048)
        file.file.seek(0)  # Reset file pointer
        
        try:
            # Try using python-magic
            mime = magic.Magic(mime=True)
            detected_mime = mime.from_buffer(file_content)
            logger.info(f"📄 File validation: {file.filename}, MIME: {detected_mime}, Ext: {file_extension}")
            
            # Check if MIME type is in allowed list
            if detected_mime in ALLOWED_MIME_TYPES:
                return True
            
            # Some leniency for text files
            if file_extension == ".txt" and "text" in detected_mime:
                return True
                
            # Fallback: allow if extension matches but MIME doesn't (common with some PDFs)
            if file_extension == ".pdf" and "pdf" in detected_mime:
                return True
                
        except Exception as e:
            logger.warning(f"MIME detection failed, falling back to extension: {e}")
            # Fallback to extension check only
            return file_extension in ALLOWED_EXTENSIONS
            
        return False
    except Exception as e:
        logger.error(f"❌ File validation failed: {e}")
        return False

def save_uploaded_file(file: UploadFile, user_id: int) -> Dict[str, Any]:
    """Save uploaded file to disk"""
    try:
        file_extension = Path(file.filename).suffix.lower()
        unique_filename = f"{uuid.uuid4().hex}{file_extension}"
        user_upload_dir = os.path.join(UPLOAD_DIR, str(user_id))
        os.makedirs(user_upload_dir, exist_ok=True)
        file_path = os.path.join(user_upload_dir, unique_filename)
        
        # Save file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        file_size = os.path.getsize(file_path)
        
        return {
            "success": True,
            "filename": unique_filename,
            "original_filename": file.filename,
            "file_path": file_path,
            "file_size": file_size,
            "file_extension": file_extension
        }
    except Exception as e:
        logger.error(f"❌ File save failed: {e}")
        return {"success": False, "error": str(e)}

# ------------------- TEXT EXTRACTION -------------------
async def extract_text_from_file(file_path: str, file_extension: str) -> Dict[str, Any]:
    """Extract text from various file formats with fallbacks"""
    try:
        text_content = ""
        metadata = {
            "page_count": 0, 
            "requires_ocr": False, 
            "ocr_confidence": 1.0, 
            "extraction_method": "direct",
            "warnings": []
        }
        
        if file_extension == ".pdf":
            result = await extract_text_from_pdf(file_path)
            if result["success"]:
                text_content = result["text_content"]
                metadata.update(result["metadata"])
            else:
                metadata["warnings"].append(f"PDF extraction failed: {result.get('error')}")
                # Fallback: return empty content with error flag
                return {"success": False, "error": result.get("error"), "metadata": metadata}
                
        elif file_extension == ".docx":
            result = await extract_text_from_docx(file_path)
            if result["success"]:
                text_content = result["text_content"]
                metadata.update(result["metadata"])
            else:
                metadata["warnings"].append(f"DOCX extraction failed: {result.get('error')}")
                return {"success": False, "error": result.get("error"), "metadata": metadata}
                
        elif file_extension == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text_content = f.read()
                metadata["extraction_method"] = "direct_read"
                metadata["word_count"] = len(text_content.split())
            except UnicodeDecodeError:
                # Try other encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        with open(file_path, "r", encoding=encoding) as f:
                            text_content = f.read()
                        metadata["extraction_method"] = f"direct_read_{encoding}"
                        metadata["word_count"] = len(text_content.split())
                        break
                    except:
                        continue
        
        # Check if text extraction was successful
        if not text_content or len(text_content.strip()) < 10:
            metadata["requires_ocr"] = True
            metadata["warnings"].append("Low text content - might require OCR")
            logger.warning(f"⚠️ Low text content: {file_path}")
        
        if "word_count" not in metadata:
            metadata["word_count"] = len(text_content.split()) if text_content else 0
        
        return {
            "success": True, 
            "text_content": text_content, 
            "metadata": metadata
        }
    except Exception as e:
        logger.error(f"❌ Text extraction failed: {e}")
        return {"success": False, "error": str(e), "metadata": {"warnings": [str(e)]}}

async def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
    """Extract text from PDF with PyPDF2 fallback"""
    try:
        text_content = ""
        metadata = {
            "page_count": 0, 
            "requires_ocr": False, 
            "ocr_confidence": 1.0, 
            "extraction_method": "pypdf2",
            "warnings": []
        }
        
        # Try PyPDF2 first
        try:
            import PyPDF2
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["page_count"] = len(pdf_reader.pages)
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text_content += f"\n--- Page {i + 1} ---\n{page_text}"
                    except:
                        text_content += f"\n--- Page {i + 1} ---\n[Text extraction failed for this page]"
        except ImportError:
            metadata["warnings"].append("PyPDF2 not installed")
            logger.error("PyPDF2 is not installed. Please run: pip install PyPDF2")
            return {"success": False, "error": "PyPDF2 not installed", "metadata": metadata}
        except Exception as e:
            metadata["warnings"].append(f"PyPDF2 extraction error: {str(e)}")
        
        # Check if we got meaningful text
        if len(text_content.strip()) < 100 and metadata["page_count"] > 0:
            metadata["requires_ocr"] = True
            metadata["extraction_method"] = "requires_ocr"
            metadata["warnings"].append("PDF may require OCR for better text extraction")
            
            # Try pdf2image + pytesseract if available
            try:
                from pdf2image import convert_from_path
                import pytesseract
                
                images = convert_from_path(file_path, first_page=1, last_page=min(3, metadata["page_count"]))
                ocr_text = ""
                for i, image in enumerate(images):
                    try:
                        ocr_text += f"\n--- Page {i + 1} (OCR) ---\n{pytesseract.image_to_string(image)}"
                    except:
                        ocr_text += f"\n--- Page {i + 1} (OCR Failed) ---\n"
                
                if len(ocr_text.strip()) > len(text_content.strip()):
                    text_content = ocr_text
                    metadata["ocr_confidence"] = 0.85
                    metadata["extraction_method"] = "ocr"
                    metadata["warnings"].append("OCR used for text extraction")
                    
            except ImportError:
                metadata["warnings"].append("OCR tools not available. Install: pip install pdf2image pytesseract")
            except Exception as e:
                metadata["warnings"].append(f"OCR attempt failed: {str(e)}")
        
        metadata["word_count"] = len(text_content.split())
        return {"success": True, "text_content": text_content, "metadata": metadata}
        
    except Exception as e:
        logger.error(f"❌ PDF extraction failed: {e}")
        return {"success": False, "error": str(e), "metadata": {"warnings": [str(e)]}}

async def extract_text_from_docx(file_path: str) -> Dict[str, Any]:
    """Extract text from DOCX files"""
    try:
        text_content = ""
        metadata = {
            "extraction_method": "docx_parser",
            "warnings": []
        }
        
        try:
            from docx import Document
            doc = Document(file_path)
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    text_content += "\t".join(row_text) + "\n"
            
        except ImportError:
            metadata["warnings"].append("python-docx not installed")
            logger.error("python-docx is not installed. Please run: pip install python-docx")
            return {"success": False, "error": "python-docx not installed", "metadata": metadata}
        except Exception as e:
            metadata["warnings"].append(f"DOCX parsing error: {str(e)}")
            return {"success": False, "error": str(e), "metadata": metadata}
        
        metadata["word_count"] = len(text_content.split())
        return {"success": True, "text_content": text_content, "metadata": metadata}
        
    except Exception as e:
        logger.error(f"❌ DOCX extraction failed: {e}")
        return {"success": False, "error": str(e), "metadata": {"warnings": [str(e)]}}

# ------------------- BACKGROUND PROCESSING -------------------
async def process_document_async(document_id: int, user_id: int, db: Session):
    """Process document in background"""
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            logger.error(f"Document {document_id} not found")
            return
        
        logger.info(f"Starting processing for document {document_id}")
        document.status = DocumentStatus.PROCESSING
        db.commit()
        
        # Extract text
        extraction_result = await extract_text_from_file(document.file_path, document.file_type)
        
        if not extraction_result["success"]:
            document.status = DocumentStatus.FAILED
            document.error_message = extraction_result.get("error", "Text extraction failed")
            db.commit()
            logger.error(f"Text extraction failed for document {document_id}")
            return
        
        text_content = extraction_result["text_content"]
        extraction_metadata = extraction_result["metadata"]
        
        # Update document metadata
        document.word_count = extraction_metadata.get("word_count", 0)
        document.page_count = extraction_metadata.get("page_count", 0)
        document.requires_ocr = extraction_metadata.get("requires_ocr", False)
        document.ocr_confidence = extraction_metadata.get("ocr_confidence", 1.0)
        document.extraction_metadata = extraction_metadata
        document.processed_at = datetime.utcnow()
        
        # Prepare metadata for RAG pipeline
        document_metadata = {
            "filename": document.filename,
            "original_filename": document.original_filename,
            "document_type": document.document_type,
            "case_number": document.case_number,
            "court_name": document.court_name,
            "judge_name": document.judge_name,
            "parties_involved": document.parties_involved or [],
            "jurisdiction": document.jurisdiction,
            "legal_topics": document.legal_topics or [],
            "user_id": user_id,
            "extraction_warnings": extraction_metadata.get("warnings", [])
        }
        
        # Process with RAG pipeline if available
        try:
            rag_pipeline = get_rag_pipeline()
            processing_result = await rag_pipeline.process_legal_document(
                user_id=user_id,
                document_id=document_id,
                document_text=text_content,
                document_metadata=document_metadata
            )
            
            if processing_result.get("success", False):
                document.status = DocumentStatus.INDEXED
                document.chunk_count = processing_result.get("chunks_indexed", 0)
                document.indexed_at = datetime.utcnow()
                logger.info(f"Document {document_id} indexed successfully")
            else:
                document.status = DocumentStatus.PROCESSED  # Mark as processed but not indexed
                document.error_message = processing_result.get("error", "RAG processing failed")
                logger.warning(f"RAG processing failed for document {document_id}")
                
        except Exception as rag_error:
            # If RAG pipeline fails, mark as processed but not indexed
            document.status = DocumentStatus.PROCESSED
            document.error_message = f"RAG pipeline error: {str(rag_error)}"
            logger.error(f"RAG pipeline error for document {document_id}: {rag_error}")
        
        db.commit()
        logger.info(f"Processing completed for document {document_id}")
        
    except Exception as e:
        logger.error(f"❌ Async document processing failed: {e}")
        logger.error(traceback.format_exc())
        try:
            if 'document' in locals():
                document.status = DocumentStatus.FAILED
                document.error_message = str(e)
                db.commit()
        except:
            pass

# ------------------- UPLOAD ENDPOINT -------------------
@router.post("/upload", response_model=UploadResponse)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    document_type: Optional[DocumentType] = Form(DocumentType.OTHER),
    case_number: Optional[str] = Form(None),
    court_name: Optional[str] = Form(None),
    judge_name: Optional[str] = Form(None),
    parties_involved: Optional[str] = Form(None),
    jurisdiction: Optional[str] = Form(None),
    legal_topics: Optional[str] = Form(None),
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Upload and process a document"""
    try:
        logger.info(f"Upload request from user {current_user.id} for file: {file.filename}")
        
        # Validate file
        if not validate_file_type(file):
            return UploadResponse(
                success=False, 
                message="Invalid file type. Only PDF, DOCX, and TXT files are allowed."
            )
        
        # Check file size
        file.file.seek(0, 2)
        file_size = file.file.tell()
        file.file.seek(0)
        
        if file_size > MAX_FILE_SIZE:
            return UploadResponse(
                success=False, 
                message=f"File too large. Maximum size is {MAX_FILE_SIZE//(1024*1024)}MB."
            )
        
        # Save file
        save_result = save_uploaded_file(file, current_user.id)
        if not save_result["success"]:
            return UploadResponse(
                success=False, 
                message=f"Failed to save file: {save_result.get('error', 'Unknown error')}"
            )
        
        # Parse JSON fields
        parties_list = []
        topics_list = []
        
        if parties_involved:
            try:
                parties_list = json.loads(parties_involved)
                if not isinstance(parties_list, list):
                    parties_list = [parties_list]
            except json.JSONDecodeError:
                parties_list = [p.strip() for p in parties_involved.split(",") if p.strip()]
            except Exception:
                parties_list = [parties_involved]
        
        if legal_topics:
            try:
                topics_list = json.loads(legal_topics)
                if not isinstance(topics_list, list):
                    topics_list = [topics_list]
            except json.JSONDecodeError:
                topics_list = [t.strip() for t in legal_topics.split(",") if t.strip()]
            except Exception:
                topics_list = [legal_topics]
        
        # Create document record
        file_name_stem = Path(save_result["original_filename"]).stem
        document_title = title or file_name_stem
        
        new_document = Document(
            user_id=current_user.id,
            filename=save_result["filename"],
            original_filename=save_result["original_filename"],
            file_path=save_result["file_path"],
            file_size=save_result["file_size"],
            file_type=save_result["file_extension"][1:],  # Remove leading dot
            title=document_title,
            description=description,
            document_type=document_type or DocumentType.OTHER,
            case_number=case_number,
            court_name=court_name,
            judge_name=judge_name,
            parties_involved=parties_list,
            jurisdiction=jurisdiction,
            legal_topics=topics_list,
            status=DocumentStatus.UPLOADED,
            uploaded_at=datetime.utcnow()
        )
        
        db.add(new_document)
        db.commit()
        db.refresh(new_document)
        
        # Start background processing
        background_tasks.add_task(process_document_async, new_document.id, current_user.id, db)
        
        logger.info(f"Document {new_document.id} uploaded successfully, processing started")
        
        return UploadResponse(
            success=True,
            message="Document uploaded successfully. Processing started in background.",
            document_id=new_document.id,
            filename=new_document.original_filename,
            file_size=new_document.file_size,
            processing_status=new_document.status
        )
        
    except Exception as e:
        db.rollback()
        logger.error(f"❌ Document upload failed: {e}")
        logger.error(traceback.format_exc())
        return UploadResponse(
            success=False, 
            message="Document upload failed. Please try again."
        )

# ------------------- UPDATE DOCUMENT ENDPOINT -------------------
@router.put("/documents/{document_id}")
async def update_document(
    document_id: int,
    document_update: DocumentUpdateModel,
    current_user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db)
):
    """Update document metadata"""
    try:
        document = db.query(Document).filter(
            Document.id == document_id, 
            Document.user_id == current_user.id
        ).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Convert Pydantic model to dict, excluding unset fields
        update_data = document_update.dict(exclude_unset=True)
        
        # Update document fields
        for field, value in update_data.items():
            if hasattr(document, field):
                setattr(document, field, value)
        
        document.updated_at = datetime.utcnow()
        db.commit()
        db.refresh(document)
        
        return {
            "success": True, 
            "message": "Document updated successfully", 
            "document": DocumentResponse.from_orm(document)
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"❌ Document update failed: {e}")
        raise HTTPException(status_code=500, detail="Failed to update document")